"""Corporate Word document generator for QAD Migration & Modernisation Plan.

Design principles:
- Only render sections/tables where AI returned real data
- Skip entire sections if data is missing or placeholder
- No TOC
- Compact spacing
"""
from __future__ import annotations

import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOWNLOADS_DIR = Path("app/static/downloads")

C = {
    "navyDark": RGBColor(0x1F, 0x4E, 0x79),
    "navyMid":  RGBColor(0x2E, 0x75, 0xB6),
    "navyLight":RGBColor(0xBD, 0xD7, 0xEE),
    "amber":    RGBColor(0xC5, 0x5A, 0x11),
    "green":    RGBColor(0x37, 0x56, 0x23),
    "red":      RGBColor(0xC0, 0x00, 0x00),
    "purple":   RGBColor(0x70, 0x30, 0xA0),
    "textDark": RGBColor(0x00, 0x00, 0x00),
    "textMid":  RGBColor(0x40, 0x40, 0x40),
    "white":    RGBColor(0xFF, 0xFF, 0xFF),
}
FILL = {
    "navyDark":  "1F4E79",
    "navyMid":   "2E75B6",
    "navyLight": "BDD7EE",
    "rowAlt":    "DEEAF1",
    "white":     "FFFFFF",
}
FONT = "Calibri"
ALT = [FILL["white"], FILL["rowAlt"]]

ACTION_COLORS = {
    "Carry Forward":       C["green"],
    "Adapt":               C["purple"],
    "Replace":             C["navyMid"],
    "Replace with Standard": C["navyMid"],
    "Decommission":        C["red"],
}
RATING_COLORS = {"HIGH": C["red"], "MEDIUM": C["amber"], "LOW": C["green"]}


def _has(value: Any) -> bool:
    if value is None:
        return False
    if isinstance(value, (list, dict)):
        return len(value) > 0
    if isinstance(value, str):
        v = value.strip()
        return bool(v) and not (v.startswith("[") and v.endswith("]"))
    return bool(value)


def _shading(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr(); tcB = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}"); el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4"); el.set(qn("w:color"), "B8CCE4")
        tcB.append(el)
    tcPr.append(tcB)


def _margins(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr(); tcM = OxmlElement("w:tcMar")
    for side, val in [("top", 80), ("bottom", 80), ("left", 140), ("right", 140)]:
        el = OxmlElement(f"w:{side}"); el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa"); tcM.append(el)
    tcPr.append(tcM)


def _cell(cell, text: str, fill: str = FILL["white"], bold: bool = False,
          italic: bool = False, color: RGBColor | None = None) -> None:
    _shading(cell, fill); _borders(cell); _margins(cell)
    p = cell.paragraphs[0]; p.clear()
    run = p.add_run(str(text) if text is not None else "")
    run.font.name = FONT; run.font.size = Pt(9); run.font.bold = bold; run.font.italic = italic
    run.font.color.rgb = color or C["textDark"]


def _hdr_cell(cell, text: str) -> None:
    _shading(cell, FILL["navyLight"]); _borders(cell); _margins(cell)
    p = cell.paragraphs[0]; p.clear()
    run = p.add_run(text); run.font.name = FONT; run.font.size = Pt(9); run.font.bold = True; run.font.color.rgb = C["navyDark"]


def _add_page_numbers(doc: Document, current_version: str, target_version: str) -> None:
    for section in doc.sections:
        footer = section.footer; footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear(); para.paragraph_format.space_before = Pt(4)
        pPr = para._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); top_el = OxmlElement("w:top")
        top_el.set(qn("w:val"), "single"); top_el.set(qn("w:sz"), "6"); top_el.set(qn("w:color"), "2E75B6"); top_el.set(qn("w:space"), "4")
        pBdr.append(top_el); pPr.append(pBdr)
        r1 = para.add_run(f"CONFIDENTIAL  ·  Mitra Central QAD-Zone  ·  AI-Generated Draft  ·  {current_version} → {target_version}    Page ")
        r1.font.name = FONT; r1.font.size = Pt(7); r1.font.color.rgb = C["amber"]
        fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
        instr1 = OxmlElement("w:instrText"); instr1.text = "PAGE"
        fldChar1e = OxmlElement("w:fldChar"); fldChar1e.set(qn("w:fldCharType"), "end")
        page_run_el = OxmlElement("w:r"); page_run_el.append(fldChar1); page_run_el.append(instr1); page_run_el.append(fldChar1e)
        para._p.append(page_run_el)


def _h1(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = FONT; run.font.size = Pt(18); run.font.bold = True; run.font.color.rgb = C["navyDark"]
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8"); bottom.set(qn("w:color"), "2E75B6"); bottom.set(qn("w:space"), "4")
    pBdr.append(bottom); pPr.append(pBdr)
    pbr = OxmlElement("w:pageBreakBefore"); pbr.set(qn("w:val"), "true"); pPr.append(pbr)
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)


def _h1_no_break(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = FONT; run.font.size = Pt(18); run.font.bold = True; run.font.color.rgb = C["navyDark"]
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8"); bottom.set(qn("w:color"), "2E75B6"); bottom.set(qn("w:space"), "4")
    pBdr.append(bottom); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(6)


def _h2(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = FONT; run.font.size = Pt(14); run.font.bold = True; run.font.color.rgb = C["navyMid"]
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)


def _h3(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.name = FONT; run.font.size = Pt(12); run.font.bold = True; run.font.color.rgb = C["textDark"]
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)


def _body(doc: Document, text: str, italic: bool = False, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text); run.font.name = FONT; run.font.size = Pt(10)
    run.font.italic = italic; run.font.color.rgb = color or C["textDark"]


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text); run.font.name = FONT; run.font.size = Pt(10); run.font.color.rgb = C["textDark"]


def _numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text); run.font.name = FONT; run.font.size = Pt(10); run.font.color.rgb = C["textDark"]


def _data_table(doc: Document, headers: list[str], rows: list[list]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=0, cols=len(headers)); table.style = "Table Grid"
    hrow = table.add_row()
    for i, h in enumerate(headers):
        _hdr_cell(hrow.cells[i], h)
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        for ci, val in enumerate(row_data):
            _cell(row.cells[ci], str(val) if val is not None else "", fill=ALT[ri % 2])


def _kv_table(doc: Document, rows: list[tuple]) -> None:
    table = doc.add_table(rows=0, cols=2); table.style = "Table Grid"
    for label, value, *extra in rows:
        color = extra[0] if extra else None
        row = table.add_row()
        _hdr_cell(row.cells[0], label)
        _cell(row.cells[1], str(value) if value else "", fill=ALT[0], bold=bool(color), color=color)


# ── Sections ──────────────────────────────────────────────────────────────────

def _cover(doc: Document, current_version: str, target_version: str, generated_on: str) -> None:
    for _ in range(4):
        sp = doc.add_paragraph(); sp.paragraph_format.space_before = Pt(5); sp.paragraph_format.space_after = Pt(0)

    p1 = doc.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0); p1.paragraph_format.space_after = Pt(10)
    r1 = p1.add_run("MITRA CENTRAL  ·  QAD-ZONE  ·  MODERNISATION")
    r1.font.name = FONT; r1.font.size = Pt(9); r1.font.bold = True; r1.font.color.rgb = C["navyMid"]

    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(6)
    pPr2 = p2._p.get_or_add_pPr(); pBdr2 = OxmlElement("w:pBdr")
    for side in ["top", "bottom"]:
        el = OxmlElement(f"w:{side}"); el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "16"); el.set(qn("w:color"), "2E75B6"); el.set(qn("w:space"), "8")
        pBdr2.append(el)
    pPr2.append(pBdr2)
    r2 = p2.add_run("QAD ERP CUSTOMISATION"); r2.font.name = FONT; r2.font.size = Pt(28); r2.font.bold = True; r2.font.color.rgb = C["navyDark"]

    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0); p3.paragraph_format.space_after = Pt(20)
    r3 = p3.add_run("MIGRATION & MODERNISATION PLAN"); r3.font.name = FONT; r3.font.size = Pt(20); r3.font.bold = True; r3.font.color.rgb = C["navyDark"]

    p4 = doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(0); p4.paragraph_format.space_after = Pt(20)
    pPr4 = p4._p.get_or_add_pPr(); shd4 = OxmlElement("w:shd")
    shd4.set(qn("w:val"), "clear"); shd4.set(qn("w:color"), "auto"); shd4.set(qn("w:fill"), FILL["navyLight"])
    pPr4.append(shd4)
    r4 = p4.add_run(f"{current_version}  →  {target_version}"); r4.font.name = FONT; r4.font.size = Pt(14); r4.font.bold = True; r4.font.color.rgb = C["navyDark"]

    for label, value, bold_val, color in [
        ("Document Date:", generated_on, False, C["textDark"]),
        ("Version:", "1.0 — AI-Generated Draft", False, C["amber"]),
    ]:
        pl = doc.add_paragraph(); pl.alignment = WD_ALIGN_PARAGRAPH.CENTER; pl.paragraph_format.space_before = Pt(0); pl.paragraph_format.space_after = Pt(4)
        rl = pl.add_run(f"{label}  "); rl.font.name = FONT; rl.font.size = Pt(10); rl.font.bold = True; rl.font.color.rgb = C["textMid"]
        rv = pl.add_run(value); rv.font.name = FONT; rv.font.size = Pt(10); rv.font.bold = bold_val; rv.font.color.rgb = color

    doc.add_page_break()


def _doc_control(doc: Document, current_version: str, target_version: str,
                 module_count: int, generated_on: str) -> None:
    _h1_no_break(doc, "Document Control")
    _h2(doc, "Document Information")
    rows = [
        ("Document Title",           "QAD ERP Customisation Migration & Modernisation Plan"),
        ("Source ERP Version",       current_version),
        ("Target ERP Version",       target_version),
        ("Number of Custom Modules", str(module_count) if module_count else "TBD"),
        ("Date Generated",           generated_on),
        ("Confidentiality",          "Confidential — For Authorised Recipients Only"),
    ]
    _data_table(doc, ["Field", "Detail"], [[l, v] for l, v in rows])
    _body(doc, "DISCLAIMER: This document has been generated by the Mitra Central QAD-Zone AI engine. All content should be reviewed by a qualified QAD implementation professional before use in project planning or budgeting. Effort estimates are indicative only.",
          italic=True, color=C["amber"])
    doc.add_page_break()


def _exec_summary(doc: Document, data: dict) -> None:
    _h1(doc, "Executive Summary")
    exec_sum = data.get("executive_summary")
    if _has(exec_sum):
        _body(doc, exec_sum)

    metrics = data.get("summary_metrics", {})
    metric_rows = [
        ("Total Custom Modules Analysed",        metrics.get("total_modules")),
        ("Recommended: Carry Forward",           metrics.get("carry_forward")),
        ("Recommended: Adapt / Re-engineer",     metrics.get("adapt")),
        ("Recommended: Replace with Standard",   metrics.get("replace")),
        ("Recommended: Decommission",            metrics.get("decommission")),
        ("Total Estimated Migration Effort",     metrics.get("total_effort")),
        ("Recommended Migration Approach",       metrics.get("approach")),
        ("Indicative Project Duration",          metrics.get("duration")),
        ("Overall Risk Rating",                  metrics.get("risk_rating")),
    ]
    valid_metrics = [[l, v] for l, v in metric_rows if _has(v)]
    if valid_metrics:
        _h2(doc, "Key Findings at a Glance")
        _data_table(doc, ["Metric", "Value"], valid_metrics)

    strategic = data.get("strategic_recommendation")
    if _has(strategic):
        _h2(doc, "Strategic Recommendation"); _body(doc, strategic)


def _introduction(doc: Document, data: dict, current_version: str, target_version: str) -> None:
    _h1(doc, "Introduction")
    _h2(doc, "Purpose of this Document")
    _body(doc, f"This Migration & Modernisation Plan supports planning the upgrade from {current_version} to {target_version}.")

    scope = data.get("scope")
    if _has(scope):
        _h2(doc, "Scope"); _body(doc, scope)

    _h2(doc, "How This Document Was Generated")
    for step in [
        "Static analysis of all custom source files to identify module names, entry points, and QAD API dependencies.",
        f"Web research to identify QAD product changes between {current_version} and {target_version}.",
        "LLM-based reasoning to determine the recommended disposition and effort estimate for each module.",
        "Aggregation of findings into this structured report.",
    ]:
        _numbered(doc, step)
    _body(doc, "NOTE: This is an AI-generated first draft. All findings must be reviewed by a qualified QAD architect before use in project planning.", italic=True, color=C["amber"])


def _system_overview(doc: Document, data: dict, current_version: str, target_version: str) -> None:
    src = data.get("source_system", {})
    tgt = data.get("target_system", {})
    arch_changes = data.get("architectural_changes")
    deprecated   = data.get("deprecated_features")

    if not src and not tgt and not _has(arch_changes) and not _has(deprecated):
        return

    _h1(doc, "Source and Target System Overview")

    if src:
        src_rows = [[k, v] for k, v in [
            ("Product Name",           src.get("product_name", current_version)),
            ("Architecture",           src.get("architecture")),
            ("Database",               src.get("database")),
            ("UI Framework",           src.get("ui_framework")),
            ("Custom Code Base Size",  src.get("code_size")),
            ("Key Business Areas",     src.get("business_areas")),
            ("Current Support Status", src.get("support_status")),
        ] if _has(v)]
        if src_rows:
            _h2(doc, f"Source System: {current_version}")
            _data_table(doc, ["Attribute", "Detail"], src_rows)

    if tgt:
        tgt_rows = [[k, v] for k, v in [
            ("Product Name",         tgt.get("product_name", target_version)),
            ("Architecture",         tgt.get("architecture")),
            ("Database",             tgt.get("database")),
            ("UI Framework",         tgt.get("ui_framework")),
            ("Custom Code Approach", tgt.get("custom_code_approach")),
            ("Key New Capabilities", tgt.get("key_capabilities")),
            ("Deployment Model",     tgt.get("deployment")),
        ] if _has(v)]
        if tgt_rows:
            _h2(doc, f"Target System: {target_version}")
            _data_table(doc, ["Attribute", "Detail"], tgt_rows)

    if _has(arch_changes):
        _h2(doc, "Key Architectural Changes Between Versions"); _body(doc, arch_changes)

    if _has(deprecated):
        rows = [[d.get("item",""), d.get("version",""), d.get("replacement","")]
                for d in deprecated if isinstance(d, dict) and _has(d.get("item"))]
        if rows:
            _h2(doc, "Deprecated APIs and Features")
            _data_table(doc, ["Deprecated Item", "Deprecated In Version", "Recommended Replacement / Action"], rows)


def _inventory(doc: Document, data: dict) -> None:
    modules = data.get("modules", [])
    if not modules:
        return

    _h1(doc, "Customisation Inventory and Summary")
    _body(doc, "The following table provides a complete inventory of all custom Progress 4GL modules identified in the source environment.")

    _h2(doc, "Master Customisation Register")
    table = doc.add_table(rows=0, cols=8); table.style = "Table Grid"
    hrow = table.add_row()
    for i, h in enumerate(["#", "Module Name", "Business Area", "Type", "Files", "Complexity", "Action", "Priority"]):
        _hdr_cell(hrow.cells[i], h)

    action_color_map = {"Low": C["green"], "Medium": C["amber"], "High": C["red"], "Very High": C["red"]}
    for idx, mod in enumerate(modules):
        if not isinstance(mod, dict):
            continue
        action = mod.get("action", "")
        complexity = mod.get("complexity", "")
        row = table.add_row()
        vals = [str(idx+1), mod.get("name",""), mod.get("business_area",""), mod.get("type",""),
                str(mod.get("files","")), complexity, action, mod.get("priority","")]
        for ci, val in enumerate(vals):
            _shading(row.cells[ci], ALT[idx % 2]); _borders(row.cells[ci]); _margins(row.cells[ci])
            p = row.cells[ci].paragraphs[0]; p.clear()
            run = p.add_run(val); run.font.name = FONT; run.font.size = Pt(9); run.font.bold = ci in (1, 6)
            if ci == 5:
                run.font.color.rgb = action_color_map.get(complexity, C["textDark"])
            elif ci == 6:
                run.font.color.rgb = ACTION_COLORS.get(action, C["textDark"])
            else:
                run.font.color.rgb = C["textDark"]

    # Summary by action
    action_summary = data.get("action_summary", {})
    if action_summary:
        summary_rows = []
        for action, count_key, pct_key in [
            ("Carry Forward",         "carry_forward_count",  "carry_forward_pct"),
            ("Adapt",                 "adapt_count",          "adapt_pct"),
            ("Replace with Standard", "replace_count",        "replace_pct"),
            ("Decommission",          "decommission_count",   "decommission_pct"),
        ]:
            cnt = action_summary.get(count_key, "")
            pct = action_summary.get(pct_key, "")
            if _has(cnt):
                summary_rows.append([action, str(cnt), str(pct)])
        if summary_rows:
            _h2(doc, "Summary by Action Type")
            _data_table(doc, ["Action", "Module Count", "% of Total"], summary_rows)


def _per_module(doc: Document, data: dict, current_version: str, target_version: str) -> None:
    modules = data.get("modules", [])
    if not modules:
        return

    _h1(doc, "Per-Module Detailed Analysis")

    for idx, mod in enumerate(modules, 1):
        if not isinstance(mod, dict):
            continue
        name   = mod.get("name", f"Module {idx}")
        action = mod.get("action", "")

        _h2(doc, f"Module {idx}: {name}")

        # Overview table — only populated fields
        overview_rows = [[k, v] for k, v in [
            ("Business Area",              mod.get("business_area")),
            ("Module Type",                mod.get("type")),
            ("Primary Source Files",       mod.get("files_list")),
            ("Total Files",                str(mod.get("files","")) if mod.get("files") else None),
            ("Lines of Code (approx.)",    mod.get("loc")),
            ("Code Complexity",            mod.get("complexity")),
            ("QAD API Dependencies",       mod.get("api_deps")),
            ("External Integrations",      mod.get("integrations")),
            ("Last Known Modification",    mod.get("last_modified")),
        ] if _has(v)]
        if overview_rows:
            _h3(doc, "Module Overview")
            _data_table(doc, ["Attribute", "Detail"], overview_rows)

        if _has(mod.get("business_purpose")):
            _h3(doc, "Business Purpose"); _body(doc, mod["business_purpose"])

        if _has(mod.get("technical_arch")):
            _h3(doc, "Technical Architecture"); _body(doc, mod["technical_arch"])

        if _has(mod.get("version_impact")):
            _h3(doc, f"Impact Assessment for {target_version}"); _body(doc, mod["version_impact"])

        # Decision table
        action_color = ACTION_COLORS.get(action)
        decision_rows = [(k, v, c) for k, v, c in [
            ("Recommended Action",            action,                         action_color),
            ("Rationale",                     mod.get("rationale"),           None),
            ("Standard QAD Feature (Replace)",mod.get("standard_feature"),    None),
            ("Effort Estimate",               mod.get("effort"),              None),
            ("Priority",                      mod.get("priority"),            None),
            ("Dependencies",                  mod.get("dependencies"),        None),
            ("Risk if Not Migrated",          mod.get("risk_if_not_migrated"),C["red"]),
        ] if _has(v)]
        if decision_rows:
            _h3(doc, "Recommended Action")
            table = doc.add_table(rows=0, cols=2); table.style = "Table Grid"
            for label, value, color in decision_rows:
                row = table.add_row()
                _hdr_cell(row.cells[0], label)
                _cell(row.cells[1], str(value), bold=(label == "Recommended Action"), color=color)

        wbs = mod.get("work_breakdown", [])
        if _has(wbs):
            wbs_rows = [[str(i+1), t.get("task",""), t.get("role",""), t.get("effort",""), t.get("notes","")]
                        for i, t in enumerate(wbs) if isinstance(t, dict) and _has(t.get("task"))]
            if wbs_rows:
                _h3(doc, "Migration Work Breakdown")
                _data_table(doc, ["#", "Task Description", "Role", "Effort", "Notes / Dependencies"], wbs_rows)

        if _has(mod.get("testing_requirements")):
            _h3(doc, "Testing Requirements"); _body(doc, mod["testing_requirements"])


def _gap_analysis(doc: Document, data: dict, target_version: str) -> None:
    gaps     = data.get("functional_gaps")
    gap_intro = data.get("gap_intro")
    new_feat  = data.get("new_features")
    if not any(_has(x) for x in [gaps, gap_intro, new_feat]):
        return

    _h1(doc, "Gap Analysis")
    if _has(gap_intro):
        _body(doc, gap_intro)

    if _has(gaps):
        gap_status_colors = {"No Gap": C["green"], "Partial Gap": C["amber"], "Full Gap Remains": C["red"]}
        table = doc.add_table(rows=0, cols=5); table.style = "Table Grid"
        hrow = table.add_row()
        for i, h in enumerate(["Business Capability", "Current Solution", f"{target_version} Standard Feature", "Gap Status", "Recommended Resolution"]):
            _hdr_cell(hrow.cells[i], h)
        for ri, gap in enumerate(gaps):
            if not isinstance(gap, dict):
                continue
            status = gap.get("gap_status", "")
            row = table.add_row()
            vals = [gap.get("capability",""), gap.get("current_solution",""), gap.get("target_feature",""), status, gap.get("resolution","")]
            for ci, val in enumerate(vals):
                _shading(row.cells[ci], ALT[ri % 2]); _borders(row.cells[ci]); _margins(row.cells[ci])
                p = row.cells[ci].paragraphs[0]; p.clear()
                run = p.add_run(val); run.font.name = FONT; run.font.size = Pt(9)
                run.font.bold = (ci == 3); run.font.color.rgb = gap_status_colors.get(status, C["textDark"]) if ci == 3 else C["textDark"]

    if _has(new_feat):
        _h2(doc, "New Standard Features to Leverage"); _body(doc, new_feat)


def _migration_strategy(doc: Document, data: dict, target_version: str) -> None:
    approach = data.get("migration_approach")
    phases   = data.get("phases")
    if not _has(approach) and not _has(phases):
        return

    _h1(doc, "Migration Strategy")
    if _has(approach):
        _h2(doc, "Recommended Approach"); _body(doc, approach)

    if _has(phases):
        rows = [[str(i), p.get("name",""), p.get("duration",""), p.get("modules",""), p.get("activities","")]
                for i, p in enumerate(phases) if isinstance(p, dict) and _has(p.get("name"))]
        if rows:
            _h2(doc, "Phasing Plan")
            _data_table(doc, ["Phase", "Name", "Duration", "Modules In Scope", "Key Activities"], rows)


def _effort(doc: Document, data: dict) -> None:
    modules = data.get("modules", [])
    if not modules:
        return

    _h1(doc, "Effort and Timeline Estimate")
    _body(doc, "The following estimates are indicative only, based on automated code analysis and AI reasoning. Refine through a formal scoping workshop before use in a project budget.",
          italic=True, color=C["amber"])

    rows = []
    total_all = 0.0
    for mod in modules:
        if not isinstance(mod, dict) or not _has(mod.get("name")):
            continue
        action = mod.get("action", "")
        eb = mod.get("effort_breakdown", {}) or {}
        an = str(eb.get("analysis", "1"))
        dv = str(eb.get("development", ""))
        te = str(eb.get("testing", ""))
        de = str(eb.get("deployment", ""))
        try:
            total = sum(float(x) for x in [an, dv, te, de] if x)
            total_all += total
            total_str = f"{total:.1f}"
        except Exception:
            total_str = ""
        if _has(mod.get("effort")):
            rows.append([mod.get("name",""), action, an, dv, te, de, total_str])

    if rows:
        _h2(doc, "Effort Estimate by Module")
        table = doc.add_table(rows=0, cols=7); table.style = "Table Grid"
        hrow = table.add_row()
        for i, h in enumerate(["Module Name", "Action", "Analysis (d)", "Dev (d)", "Test (d)", "Deploy (d)", "Total (d)"]):
            _hdr_cell(hrow.cells[i], h)
        for ri, row_data in enumerate(rows):
            row = table.add_row()
            action = row_data[1]
            for ci, val in enumerate(row_data):
                _shading(row.cells[ci], ALT[ri % 2]); _borders(row.cells[ci]); _margins(row.cells[ci])
                p = row.cells[ci].paragraphs[0]; p.clear()
                run = p.add_run(str(val)); run.font.name = FONT; run.font.size = Pt(9); run.font.bold = ci in (0, 6)
                run.font.color.rgb = ACTION_COLORS.get(action, C["textDark"]) if ci == 1 else C["textDark"]
        if total_all > 0:
            trow = table.add_row()
            for ci, val in enumerate(["TOTAL", "", "", "", "", "", f"~{total_all:.0f}"]):
                _shading(trow.cells[ci], FILL["navyLight"]); _borders(trow.cells[ci]); _margins(trow.cells[ci])
                p = trow.cells[ci].paragraphs[0]; p.clear()
                run = p.add_run(val); run.font.name = FONT; run.font.size = Pt(9); run.font.bold = True; run.font.color.rgb = C["navyDark"]

    timeline = data.get("timeline_narrative")
    if _has(timeline):
        _h2(doc, "Indicative Timeline"); _body(doc, timeline)


def _risk(doc: Document, data: dict) -> None:
    risks = data.get("risks", [])
    # Always include standard risks + any AI-generated ones
    standard_risks = [
        ("R01", "Key Progress 4GL developer unavailable during migration", "Delays to Adapt modules; go-live slip", "3", "5", "15", "HIGH", "Pre-book developer resource; identify backup"),
        ("R02", "QAD target version API incompatible with assumed approach", "Re-engineering effort higher than estimated", "2", "4", "8", "MEDIUM", "Validate API assumptions in sandbox before development"),
        ("R03", "Business stakeholder unavailability for UAT", "UAT sign-off delayed; go-live pushed", "3", "3", "9", "MEDIUM", "Schedule UAT windows at project kick-off"),
        ("R04", "Undocumented custom code found during analysis", "Scope creep; additional effort", "2", "3", "6", "LOW", "Run full Mitra Central scan before project freeze"),
    ]
    all_risks = list(standard_risks)
    for risk in (risks or []):
        if isinstance(risk, dict) and _has(risk.get("description")):
            all_risks.append((
                risk.get("id", f"R{len(all_risks)+1:02d}"),
                risk.get("description", ""),
                risk.get("impact", ""),
                str(risk.get("likelihood", "")),
                str(risk.get("impact_score", "")),
                str(risk.get("score", "")),
                risk.get("rating", ""),
                risk.get("mitigation", ""),
            ))

    _h1(doc, "Risk Register")
    table = doc.add_table(rows=0, cols=8); table.style = "Table Grid"
    hrow = table.add_row()
    for i, h in enumerate(["ID", "Risk Description", "Impact", "Likelihood (1–5)", "Impact (1–5)", "Score", "Rating", "Mitigation"]):
        _hdr_cell(hrow.cells[i], h)
    for ri, risk_row in enumerate(all_risks):
        rating = risk_row[6]
        color = RATING_COLORS.get(rating, C["textDark"])
        row = table.add_row()
        for ci, val in enumerate(risk_row):
            _shading(row.cells[ci], ALT[ri % 2]); _borders(row.cells[ci]); _margins(row.cells[ci])
            p = row.cells[ci].paragraphs[0]; p.clear()
            run = p.add_run(str(val)); run.font.name = FONT; run.font.size = Pt(9)
            run.font.bold = ci in (5, 6); run.font.color.rgb = color if ci in (5, 6) else C["textDark"]


def _recommendations(doc: Document, data: dict, target_version: str) -> None:
    recs      = data.get("recommendations", [])
    recs_intro = data.get("recommendations_intro")
    if not _has(recs) and not _has(recs_intro):
        return

    _h1(doc, "Recommendations and Next Steps")
    if _has(recs_intro):
        _body(doc, recs_intro)

    if _has(recs):
        rows = [[str(i+1), r.get("recommendation",""), r.get("owner",""), r.get("date","")]
                for i, r in enumerate(recs) if isinstance(r, dict) and _has(r.get("recommendation"))]
        if rows:
            _data_table(doc, ["#", "Recommendation", "Owner", "Target Date"], rows)

    _h2(doc, "Immediate Next Steps (Within 30 Days)")
    for step in [
        "Review and validate this AI-generated plan with the QAD project team and implementation partner.",
        "Conduct a code freeze on the source QAD environment.",
        f"Establish a dedicated QAD {target_version} sandbox environment for developer testing.",
        "Assign business owners to each custom module and confirm recommended actions.",
        "Schedule a risk workshop to review the risk register and agree mitigation owners.",
    ]:
        _numbered(doc, step)


def _appendix(doc: Document, data: dict, current_version: str, target_version: str) -> None:
    source_inv = data.get("source_file_inventory")
    ver_ref    = data.get("version_change_reference")
    if not _has(source_inv) and not _has(ver_ref):
        return

    _h1(doc, "Appendices")

    if _has(source_inv):
        _h2(doc, "Appendix A: Source Code File Inventory")
        rows = [[sf.get("path",""), sf.get("ext",""), str(sf.get("loc","")), sf.get("module",""), sf.get("notes","")]
                for sf in source_inv if isinstance(sf, dict) and _has(sf.get("path"))]
        if rows:
            _data_table(doc, ["File Path", "Extension", "Lines of Code", "Module", "Notes"], rows)

    if _has(ver_ref):
        _h2(doc, "Appendix B: QAD Version Change Reference")
        _body(doc, ver_ref)

    _h2(doc, "Appendix C: Glossary")
    _data_table(doc, ["Term", "Definition"], [
        ["4GL / ABL",        "Progress 4th Generation Language / OpenEdge ABL — the language used to write QAD customisations."],
        ["Carry Forward",    "Customisation compatible with target version; redeployable with minimal changes."],
        ["QCF",              "QAD Customisation Framework — the supported hook-based customisation mechanism in modern QAD."],
        ["QAD Adaptive ERP", "QAD's cloud-native ERP platform, successor to QAD Enterprise Edition."],
        ["TCO",              "Total Cost of Ownership — licence, infrastructure, and ongoing maintenance cost of custom code."],
    ])

    _h2(doc, "Appendix D: About Mitra Central QAD-Zone")
    _body(doc, "This Migration & Modernisation Plan was generated by the Mitra Central QAD-Zone Modernisation engine — an AI-powered platform designed for QAD ERP implementation partners. For queries, contact [PARTNER CONTACT DETAILS].")


# ── Main entry point ──────────────────────────────────────────────────────────

def generate_migration_document(
    current_version: str,
    target_version: str,
    data: dict,
) -> str:
    DOWNLOADS_DIR.mkdir(parents=True, exist_ok=True)
    generated_on = datetime.now().strftime("%Y-%m-%d")
    modules = data.get("modules", [])
    module_count = len(modules)

    doc = Document()
    for section in doc.sections:
        section.page_width = Pt(11906 / 20); section.page_height = Pt(16838 / 20)
        for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(section, attr, Pt(1440 / 20))

    # Header
    header = doc.sections[0].header
    h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    h_para.clear(); h_para.paragraph_format.space_after = Pt(4)
    pPr = h_para._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6"); bot.set(qn("w:color"), "2E75B6"); bot.set(qn("w:space"), "4")
    pBdr.append(bot); pPr.append(pBdr)
    r1 = h_para.add_run(f"QAD ERP Migration & Modernisation Plan")
    r1.font.name = FONT; r1.font.size = Pt(8); r1.font.color.rgb = C["textMid"]
    r2 = h_para.add_run(f"    {current_version} → {target_version}")
    r2.font.name = FONT; r2.font.size = Pt(8); r2.font.color.rgb = C["navyMid"]; r2.font.bold = True

    _add_page_numbers(doc, current_version, target_version)

    _cover(doc, current_version, target_version, generated_on)
    _doc_control(doc, current_version, target_version, module_count, generated_on)
    _exec_summary(doc, data)
    _introduction(doc, data, current_version, target_version)
    _system_overview(doc, data, current_version, target_version)
    _inventory(doc, data)
    _per_module(doc, data, current_version, target_version)
    _gap_analysis(doc, data, target_version)
    _migration_strategy(doc, data, target_version)
    _effort(doc, data)
    _risk(doc, data)
    _recommendations(doc, data, target_version)
    _appendix(doc, data, current_version, target_version)

    filename = f"{uuid.uuid4().hex[:12]}.docx"
    doc.save(str(DOWNLOADS_DIR / filename))
    return f"/static/downloads/{filename}"