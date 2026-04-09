"""Corporate Word document generator for QAD Custom Module Documentation.

Design principles:
- Only render sections where AI actually returned real data
- Skip entire sections (heading + content) if data is missing or placeholder
- No TOC (removed — requires manual Word update)
- Compact spacing, no large empty blocks
- Structure follows QAD customisation documentation best practice:
    Cover → Doc Control → Executive Summary → Module Overview →
    Technical Design → Program Documentation → Business Rules →
    Security & Testing → Glossary → AI Log
"""
from __future__ import annotations

import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOWNLOADS_DIR = Path("app/static/downloads")

C = {
    "navy":      RGBColor(0x1F, 0x38, 0x64),
    "blue":      RGBColor(0x2E, 0x75, 0xB6),
    "lightBlue": RGBColor(0xD9, 0xE2, 0xF3),
    "white":     RGBColor(0xFF, 0xFF, 0xFF),
    "darkGrey":  RGBColor(0x40, 0x40, 0x40),
    "midGrey":   RGBColor(0x76, 0x76, 0x76),
    "lightGrey": RGBColor(0xF2, 0xF2, 0xF2),
    "orange":    RGBColor(0xED, 0x7D, 0x31),
    "red":       RGBColor(0xC0, 0x00, 0x00),
}
FILL = {
    "navy":      "1F3864",
    "blue":      "2E75B6",
    "paleBlue":  "EEF3FB",
    "white":     "FFFFFF",
    "lightGrey": "F2F2F2",
}
FONT = "Calibri"
ALT = [FILL["white"], FILL["lightGrey"]]


def _has(value: Any) -> bool:
    """True only if value is real, non-empty, non-placeholder content."""
    if value is None:
        return False
    if isinstance(value, (list, dict)):
        return len(value) > 0
    if isinstance(value, str):
        v = value.strip()
        return bool(v) and not (v.startswith("[") and v.endswith("]"))
    return bool(value)


def _shading(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4"); el.set(qn("w:color"), "BFBFBF")
        tcB.append(el)
    tcPr.append(tcB)


def _margins(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcM = OxmlElement("w:tcMar")
    for side, val in [("top", 80), ("bottom", 80), ("left", 140), ("right", 140)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        tcM.append(el)
    tcPr.append(tcM)


def _cell(cell, text: str, fill: str = FILL["white"], bold: bool = False,
          italic: bool = False, color: RGBColor | None = None) -> None:
    _shading(cell, fill); _borders(cell); _margins(cell)
    p = cell.paragraphs[0]; p.clear()
    run = p.add_run(str(text))
    run.font.name = FONT; run.font.size = Pt(10)
    run.font.bold = bold; run.font.italic = italic
    run.font.color.rgb = color or C["darkGrey"]


def _hdr_cell(cell, text: str, fill: str = FILL["navy"]) -> None:
    _shading(cell, fill); _borders(cell); _margins(cell)
    p = cell.paragraphs[0]; p.clear()
    run = p.add_run(text)
    run.font.name = FONT; run.font.size = Pt(10); run.font.bold = True; run.font.color.rgb = C["white"]


def _add_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.paragraph_format.space_before = Pt(4)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top_el = OxmlElement("w:top")
        top_el.set(qn("w:val"), "single"); top_el.set(qn("w:sz"), "6")
        top_el.set(qn("w:color"), "1F3864"); top_el.set(qn("w:space"), "1")
        pBdr.append(top_el); pPr.append(pBdr)
        r1 = para.add_run("INTERNAL — IT/ERP Team Use Only    Mitra Central — Auto-Generated    Page ")
        r1.font.name = FONT; r1.font.size = Pt(8); r1.font.color.rgb = C["midGrey"]; r1.font.italic = True
        fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
        instr1 = OxmlElement("w:instrText"); instr1.text = "PAGE"
        fldChar1e = OxmlElement("w:fldChar"); fldChar1e.set(qn("w:fldCharType"), "end")
        page_run_el = OxmlElement("w:r")
        page_run_el.append(fldChar1); page_run_el.append(instr1); page_run_el.append(fldChar1e)
        para._p.append(page_run_el)


def _h1(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = FONT; run.font.size = Pt(15); run.font.bold = True; run.font.color.rgb = C["white"]
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), FILL["navy"])
    pPr.append(shd)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "180"); ind.set(qn("w:right"), "180")
    pPr.append(ind)
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)


def _h2(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = FONT; run.font.size = Pt(13); run.font.bold = True; run.font.color.rgb = C["navy"]
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4"); bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)


def _h3(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.name = FONT; run.font.size = Pt(11); run.font.bold = True; run.font.color.rgb = C["blue"]
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)


def _body(doc: Document, text: str, italic: bool = False, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT; run.font.size = Pt(10); run.font.italic = italic
    run.font.color.rgb = color or C["darkGrey"]


def _warn(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "FCE4E4")
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr"); left = OxmlElement("w:left")
    left.set(qn("w:val"), "thick"); left.set(qn("w:sz"), "10"); left.set(qn("w:space"), "6"); left.set(qn("w:color"), "C00000")
    pBdr.append(left); pPr.append(pBdr)
    r1 = p.add_run("WARNING: "); r1.font.name = FONT; r1.font.size = Pt(10); r1.font.bold = True; r1.font.color.rgb = C["red"]
    r2 = p.add_run(text); r2.font.name = FONT; r2.font.size = Pt(10); r2.font.color.rgb = C["darkGrey"]


def _code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), FILL["lightGrey"])
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr"); left = OxmlElement("w:left")
    left.set(qn("w:val"), "thick"); left.set(qn("w:sz"), "10"); left.set(qn("w:space"), "6"); left.set(qn("w:color"), "2E75B6")
    pBdr.append(left); pPr.append(pBdr)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "320"); pPr.append(ind)
    run = p.add_run(text); run.font.name = "Courier New"; run.font.size = Pt(9); run.font.color.rgb = C["darkGrey"]


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text); run.font.name = FONT; run.font.size = Pt(10); run.font.color.rgb = C["darkGrey"]


def _kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2); table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row()
        _hdr_cell(row.cells[0], label, fill=FILL["blue"])
        _cell(row.cells[1], value, fill=FILL["paleBlue"], italic=True, color=C["blue"])


def _data_table(doc: Document, headers: list[str], data_rows: list[list[str]]) -> None:
    if not data_rows:
        return
    table = doc.add_table(rows=0, cols=len(headers)); table.style = "Table Grid"
    hrow = table.add_row()
    for i, h in enumerate(headers):
        _hdr_cell(hrow.cells[i], h)
    for ri, row_data in enumerate(data_rows):
        row = table.add_row()
        for ci, val in enumerate(row_data):
            _cell(row.cells[ci], str(val) if val is not None else "", fill=ALT[ri % 2])


# ── Section builders — each checks _has() before rendering ────────────────────

def _cover(doc: Document, data: dict) -> None:
    company_name = data.get("company_name", "Yash Technologies Pvt Ltd.")
    module_name  = data.get("module_name", "")
    module_title = data.get("module_title", "")
    module_code  = data.get("module_code", "")
    qad_version  = data.get("qad_version", "")

    for line, size in [(company_name, 24), ("IT / ERP Centre of Excellence  ·  QAD Custom Modules", 10), (" ", 7)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), FILL["navy"])
        pPr.append(shd)
        r = p.add_run(line); r.font.name = FONT; r.font.size = Pt(size); r.font.bold = (size == 24); r.font.color.rgb = C["lightBlue"]

    sp = doc.add_paragraph(); sp.paragraph_format.space_before = Pt(20); sp.paragraph_format.space_after = Pt(0)

    if _has(module_name):
        pn = doc.add_paragraph(); pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pn.paragraph_format.space_before = Pt(0); pn.paragraph_format.space_after = Pt(4)
        rn = pn.add_run(module_name); rn.font.name = FONT; rn.font.size = Pt(32); rn.font.bold = True; rn.font.color.rgb = C["navy"]

    if _has(module_title) and module_title != module_name:
        pt = doc.add_paragraph(); pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt.paragraph_format.space_before = Pt(0); pt.paragraph_format.space_after = Pt(8)
        rt = pt.add_run(module_title); rt.font.name = FONT; rt.font.size = Pt(16); rt.font.color.rgb = C["blue"]

    pl = doc.add_paragraph(); pl.paragraph_format.space_before = Pt(4); pl.paragraph_format.space_after = Pt(12)
    pPr = pl._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4"); bot.set(qn("w:color"), "2E75B6"); bot.set(qn("w:space"), "4")
    pBdr.append(bot); pPr.append(pBdr)

    meta_rows = [("Document Type", "Technical Reference — Custom Module Documentation")]
    if _has(module_code):  meta_rows.append(("Module Code", module_code))
    if _has(qad_version):  meta_rows.append(("QAD Version", qad_version))
    meta_rows.append(("Generated On", datetime.now().strftime("%Y-%m-%d")))
    meta_rows.append(("Classification", "INTERNAL — IT/ERP Team Only"))
    table = doc.add_table(rows=0, cols=2); table.style = "Table Grid"
    for label, value in meta_rows:
        row = table.add_row(); _hdr_cell(row.cells[0], label, fill=FILL["blue"]); _cell(row.cells[1], value)

    sp2 = doc.add_paragraph(); sp2.paragraph_format.space_before = Pt(10)
    pa = doc.add_paragraph(); pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ra = pa.add_run("Auto-generated by Mitra Central — QAD-Zone Documentation Mode")
    ra.font.name = FONT; ra.font.size = Pt(8); ra.font.italic = True; ra.font.color.rgb = C["midGrey"]
    doc.add_page_break()


def _doc_control(doc: Document) -> None:
    _h1(doc, "1.  DOCUMENT CONTROL")
    _h2(doc, "Version History")
    _data_table(doc, ["Version", "Date", "Author", "Reviewer", "Summary of Changes"],
                [["1.0.0", datetime.now().strftime("%Y-%m-%d"), "[NAME]", "[NAME]", "Initial auto-generated documentation."]])
    _h2(doc, "Approvals")
    _data_table(doc, ["Role", "Name", "Date", "Sign-off Reference"],
                [["IT / ERP Manager", "[NAME]", "[DATE]", "[Email / Ticket Ref]"],
                 ["Business Module Owner", "[NAME]", "[DATE]", "[Email / Ticket Ref]"],
                 ["QA / Test Lead", "[NAME]", "[DATE]", "[Email / Ticket Ref]"]])
    doc.add_page_break()


def _exec_summary(doc: Document, data: dict) -> None:
    _h1(doc, "2.  EXECUTIVE SUMMARY")
    fields = [
        ("Module Code",        data.get("module_code")),
        ("Full Name",          data.get("module_full_name") or data.get("module_title")),
        ("Business Domain",    data.get("business_domain")),
        ("Primary Purpose",    data.get("primary_purpose")),
        ("Owning Department",  data.get("owning_department")),
        ("Development Status", data.get("dev_status")),
        ("Criticality",        data.get("criticality")),
        ("Total Source Files", data.get("total_files")),
    ]
    valid = [(l, v) for l, v in fields if _has(v)]
    if valid:
        _kv_table(doc, valid)

    caps = data.get("key_capabilities")
    if _has(caps):
        _h2(doc, "Key Capabilities")
        for cap in caps:
            if _has(cap):
                _bullet(doc, cap)

    scope_in  = data.get("scope_in")
    scope_out = data.get("scope_out")
    if _has(scope_in) or _has(scope_out):
        _h2(doc, "Scope")
        if _has(scope_in):
            _body(doc, "In scope:"); _body(doc, scope_in)
        if _has(scope_out):
            _body(doc, "Out of scope:"); _body(doc, scope_out)
    doc.add_page_break()


def _module_overview(doc: Document, data: dict) -> None:
    background   = data.get("background")
    process_flow = data.get("process_flow")
    user_roles   = data.get("user_roles")
    if not any(_has(x) for x in [background, process_flow, user_roles]):
        return
    _h1(doc, "3.  MODULE OVERVIEW")
    if _has(background):
        _h2(doc, "Background & Business Context"); _body(doc, background)
    if _has(process_flow):
        _h2(doc, "End-to-End Process Flow"); _body(doc, process_flow)
    if _has(user_roles):
        rows = [[r.get("role",""), r.get("responsibilities",""), r.get("domain",""), r.get("notes","")]
                for r in user_roles if isinstance(r, dict) and any(_has(r.get(k)) for k in ["role","responsibilities"])]
        if rows:
            _h2(doc, "User Roles & Stakeholders")
            _data_table(doc, ["Role / User Type", "Responsibilities", "QAD Security Domain", "Notes"], rows)
    doc.add_page_break()


def _technical_design(doc: Document, data: dict) -> None:
    arch     = data.get("architecture")
    intgs    = data.get("integrations")
    ctables  = data.get("custom_tables")
    kfields  = data.get("key_fields")
    stables  = data.get("standard_tables")
    if not any(_has(x) for x in [arch, intgs, ctables, kfields, stables]):
        return
    _h1(doc, "4.  TECHNICAL DESIGN")
    if _has(arch):
        _h2(doc, "Architecture & Integration"); _body(doc, arch)
    if _has(intgs):
        rows = [[i.get("system",""), i.get("code",""), i.get("type",""), i.get("data",""), i.get("direction","")]
                for i in intgs if isinstance(i, dict) and _has(i.get("system"))]
        if rows:
            _h2(doc, "Integration Map")
            _data_table(doc, ["System / Module", "Code", "Integration Type", "Data Exchanged", "Direction"], rows)
    if _has(ctables) or _has(kfields) or _has(stables):
        _h2(doc, "Database Objects")
        if _has(ctables):
            rows = [[t.get("name",""), t.get("df_file",""), t.get("purpose","")]
                    for t in ctables if isinstance(t, dict) and _has(t.get("name"))]
            if rows:
                _h3(doc, "Custom Tables")
                _data_table(doc, ["Table Name", "Source .df File", "Purpose"], rows)
        if _has(kfields):
            rows = [[f.get("field",""), f.get("label",""), f.get("type",""), f.get("key",""), f.get("desc","")]
                    for f in kfields if isinstance(f, dict) and _has(f.get("field"))]
            if rows:
                _h3(doc, "Key Fields Reference")
                _data_table(doc, ["Table.Field", "Label", "Type", "Key", "Description"], rows)
        if _has(stables):
            rows = [[t.get("table",""), t.get("owner",""), t.get("access",""), t.get("purpose","")]
                    for t in stables if isinstance(t, dict) and _has(t.get("table"))]
            if rows:
                _h3(doc, "Standard QAD Tables Accessed")
                _data_table(doc, ["Table", "QAD Module Owner", "Access", "Purpose / Fields Used"], rows)
    doc.add_page_break()


def _program_doc(doc: Document, data: dict) -> None:
    source_files = data.get("source_files")
    key_programs = data.get("key_programs")
    if not _has(source_files) and not _has(key_programs):
        return
    _h1(doc, "5.  PROGRAM DOCUMENTATION")
    if _has(source_files):
        rows = [[p.get("name",""), p.get("type",""), str(p.get("lines","")), p.get("purpose","")]
                for p in source_files if isinstance(p, dict) and _has(p.get("name"))]
        if rows:
            _h2(doc, "Source File Inventory")
            _data_table(doc, ["File Name", "Type", "Lines", "Purpose"], rows)
    if _has(key_programs):
        _h2(doc, "Key Program Detail")
        for idx, prog in enumerate(key_programs, 1):
            if not isinstance(prog, dict):
                continue
            _h3(doc, f"{idx}.  {prog.get('name', f'Program {idx}')}")
            meta = [(l, prog.get(k)) for l, k in [
                ("Type", "type"), ("Called By", "called_by"), ("Calls", "calls"),
                ("Tables (Write)", "tables_write"), ("Tables (Read)", "tables_read"),
                ("Include Files", "includes"),
            ] if _has(prog.get(k))]
            if meta:
                _kv_table(doc, meta)
            if _has(prog.get("logic_flow")):
                _h3(doc, "Logic Flow"); _body(doc, prog["logic_flow"])
            if _has(prog.get("code_snippet")):
                _h3(doc, "Representative Code Snippet")
                for line in prog["code_snippet"].split("\n"):
                    _code(doc, line)
            if _has(prog.get("error_handling")):
                _h3(doc, "Error Handling & Performance Notes"); _body(doc, prog["error_handling"])
    doc.add_page_break()


def _business_rules(doc: Document, data: dict) -> None:
    rules  = data.get("business_rules")
    config = data.get("config_params")
    audit  = data.get("audit_trail")
    if not any(_has(x) for x in [rules, config, audit]):
        return
    _h1(doc, "6.  BUSINESS RULES & CONFIGURATION")
    if _has(rules):
        rows = [[r.get("name",""), r.get("description",""), r.get("enforced_in",""), r.get("consequence","")]
                for r in rules if isinstance(r, dict) and _has(r.get("description"))]
        if rows:
            _h2(doc, "Core Business Rules")
            _data_table(doc, ["Rule Name", "Rule Description", "Enforced In", "Consequence"], rows)
    if _has(config):
        rows = [[c.get("param",""), c.get("stored_in",""), c.get("default",""), c.get("desc","")]
                for c in config if isinstance(c, dict) and _has(c.get("param"))]
        if rows:
            _h2(doc, "Configuration Parameters")
            _data_table(doc, ["Parameter / Setting", "Stored In", "Default", "Description"], rows)
    if _has(audit):
        _h2(doc, "Audit Trail"); _body(doc, audit)
        if any(w in audit.lower() for w in ["no audit", "not implemented", "not logged", "none"]):
            _warn(doc, "No audit trail is implemented. This is a compliance risk in regulated environments.")
    doc.add_page_break()


def _security_testing(doc: Document, data: dict) -> None:
    sec    = data.get("security_objects")
    tests  = data.get("test_cases")
    issues = data.get("known_issues")
    if not any(_has(x) for x in [sec, tests, issues]):
        return
    _h1(doc, "7.  SECURITY, TESTING & HANDOVER")
    if _has(sec):
        rows = [[s.get("object",""), s.get("type",""), s.get("role",""), s.get("notes","")]
                for s in sec if isinstance(s, dict) and _has(s.get("object"))]
        if rows:
            _h2(doc, "Access Control")
            _data_table(doc, ["Security Object", "Type", "Minimum Role / Token", "Notes"], rows)
    if _has(tests):
        rows = [[t.get("id",""), t.get("scenario",""), t.get("steps",""), t.get("expected",""), t.get("status","")]
                for t in tests if isinstance(t, dict) and _has(t.get("scenario"))]
        if rows:
            _h2(doc, "Test Scenarios")
            _data_table(doc, ["TC #", "Scenario", "Steps", "Expected Result", "Status"], rows)
    if _has(issues):
        rows = [[i.get("id",""), i.get("severity",""), i.get("description",""), i.get("workaround",""), i.get("status","")]
                for i in issues if isinstance(i, dict) and _has(i.get("description"))]
        if rows:
            _h2(doc, "Known Issues & Limitations")
            _data_table(doc, ["Issue #", "Severity", "Description", "Workaround", "Status"], rows)
    doc.add_page_break()


def _glossary(doc: Document, data: dict) -> None:
    terms = data.get("glossary_terms")
    if not _has(terms):
        return
    rows = [[t.get("term",""), t.get("definition","")]
            for t in terms if isinstance(t, dict) and _has(t.get("term")) and _has(t.get("definition"))]
    if not rows:
        return
    _h1(doc, "8.  GLOSSARY")
    _data_table(doc, ["Term / Abbreviation", "Definition"], rows)


def _ai_log(doc: Document, data: dict) -> None:
    _h1(doc, "APPENDIX — AI GENERATION LOG")
    rows = [("Generated By", "Mitra Central — QAD-Zone Documentation Mode"),
            ("Timestamp", datetime.now().strftime("%Y-%m-%dT%H:%M:%S")),
            ("LLM Model", "gpt-4o")]
    if _has(data.get("files_scanned")):
        rows.insert(2, ("Files Scanned", data["files_scanned"]))
    if _has(data.get("lines_analysed")):
        rows.insert(3, ("Lines Analysed", data["lines_analysed"]))
    _kv_table(doc, rows)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    _warn(doc, "AI-generated content must be reviewed and approved by the module owner before this document is stored as an official record.")


# ── Main entry point ──────────────────────────────────────────────────────────

def generate_document(title: str, sections: list[dict], *, subtitle: str = "Mitra Central") -> str:
    DOWNLOADS_DIR.mkdir(parents=True, exist_ok=True)

    # Extract structured data
    data: dict[str, Any] = {}
    for s in sections:
        if s.get("heading") == "structured_data" and s.get("metadata"):
            data = s["metadata"]
            break

    # Legacy fallback
    if not data:
        data = {"module_name": title, "module_title": title}
        for s in sections:
            h = (s.get("heading") or "").lower()
            content = (s.get("content") or "").strip()
            if not content or (content.startswith("[") and content.endswith("]")):
                continue
            if "background" in h or "context" in h:
                data["background"] = content
            elif "purpose" in h:
                data["primary_purpose"] = content
            elif "process flow" in h:
                data["process_flow"] = content
            elif "architecture" in h:
                data["architecture"] = content
            elif "audit" in h:
                data["audit_trail"] = content

    doc = Document()
    for section in doc.sections:
        section.page_width = Pt(11906 / 20); section.page_height = Pt(16838 / 20)
        for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(section, attr, Pt(1134 / 20))

    # Header
    header = doc.sections[0].header
    h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    h_para.clear(); h_para.paragraph_format.space_after = Pt(4)
    pPr = h_para._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6"); bot.set(qn("w:color"), "1F3864"); bot.set(qn("w:space"), "1")
    pBdr.append(bot); pPr.append(pBdr)
    company = data.get("company_name", "Yash Technologies Pvt Ltd.")
    mod_name = data.get("module_name", title)
    r1 = h_para.add_run(f"{company}  |  QAD Custom Module — {mod_name}")
    r1.font.name = FONT; r1.font.size = Pt(9); r1.font.color.rgb = C["midGrey"]
    r2 = h_para.add_run("    v1.0")
    r2.font.name = FONT; r2.font.size = Pt(9); r2.font.color.rgb = C["blue"]

    _add_page_numbers(doc)
    _cover(doc, data)
    _doc_control(doc)
    _exec_summary(doc, data)
    _module_overview(doc, data)
    _technical_design(doc, data)
    _program_doc(doc, data)
    _business_rules(doc, data)
    _security_testing(doc, data)
    _glossary(doc, data)
    _ai_log(doc, data)

    filename = f"{uuid.uuid4().hex[:12]}.docx"
    doc.save(str(DOWNLOADS_DIR / filename))
    return f"/static/downloads/{filename}"